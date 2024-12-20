import shutil
import docx
from copy import deepcopy


class DocumentGenerator:
    def __init__(
        self, template_path, output_path, employee_name, approver, reporting_month, year
    ):
        shutil.copyfile(template_path, output_path)
        self.doc = docx.Document(output_path)
        self.output_path = output_path
        self.employee_name = employee_name
        self.approver = approver
        self.reporting_month = reporting_month
        self.year = year

    def setup_employee_data(self):
        table = self.doc.tables[0]
        table.cell(0, 1).text = self.employee_name
        table.cell(1, 1).text = f"{self.reporting_month} {self.year}"
        table.cell(2, 1).text = self.approver

    def add_hyperlink(self, paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)
        new_run = docx.oxml.shared.OxmlElement("w:r")
        rPr = docx.oxml.shared.OxmlElement("w:rPr")
        c = docx.oxml.shared.OxmlElement("w:color")
        c.set(docx.oxml.shared.qn("w:val"), "0563c1")
        rPr.append(c)
        u = docx.oxml.shared.OxmlElement("w:u")
        u.set(docx.oxml.shared.qn("w:val"), "single")
        rPr.append(u)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def add_repository_data(self, url, commits):
        paragraph = self.doc.add_paragraph()
        template_table = deepcopy(self.doc.tables[1])
        paragraph._p.addprevious(template_table._tbl)
        new_table = self.doc.tables[-1]
        url_paragraph = new_table.cell(0, 1).paragraphs[0]
        self.add_hyperlink(url_paragraph, url, url)
        new_table.cell(1, 1).text = commits

    def finalize(self):
        self._delete_template_table()
        self.doc.save(self.output_path)

    def _delete_template_table(self):
        table = self.doc.tables[1]
        table._element.getparent().remove(table._element)
