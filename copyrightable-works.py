import calendar
import docx
import os
import shutil
import subprocess
from dotenv import load_dotenv
from datetime import datetime
from pathlib import Path
from copy import deepcopy

load_dotenv();


name = os.environ.get('NAME')
approver = os.environ.get('APPROVER')
author = os.environ.get('AUTHOR')
author_email = os.environ.get('AUTHOR_EMAIL')
repos = os.environ.get('REPOS').split(',')
filename = f'{datetime.today().year}.{str(datetime.today().month).zfill(2)}_{name.replace(" ", "_")}.docx'


Path("./dist").mkdir(parents=True, exist_ok=True)

shutil.copyfile('./assets/template.docx', f'./dist/{filename}')

doc = docx.Document(f'./dist/{filename}')

reporting_month = calendar.month_name[datetime.today().month]

employee_data_table = doc.tables[0]
employee_data_table.cell(0, 1).text = name
employee_data_table.cell(1, 1).text = f'{reporting_month.capitalize()} {datetime.today().year}'
employee_data_table.cell(2, 1).text = approver


def delete_table(table):
        table._element.getparent().remove(table._element)


def run_command_in_repo(command, repo):
    process  = subprocess.Popen(command, stdout = subprocess.PIPE, shell=True, universal_newlines = True, stderr = subprocess.PIPE, text=True, cwd=repo)
    out, err = process.communicate()
    if err: print(err)
    return out


def get_commits_command():
    month_start_date = f'01.{str(datetime.today().month).zfill(2)}.{datetime.today().year}'
    return f'git log origin/main --author={author_email} --author={author} --since=\'{month_start_date}\' --reverse'


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0563c1')
    rPr.append(c)
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


for repo in repos:
    remote_url = run_command_in_repo('git config --get remote.origin.url', repo).strip().replace('.git', '').replace('git@github.com:', 'https://github.com/')
    commits = run_command_in_repo(get_commits_command(), repo)
    if (commits): print('Found commits: ', remote_url)
    if commits:
        paragraph = doc.add_paragraph()
        template_table = deepcopy(doc.tables[1])
        paragraph._p.addprevious(template_table._tbl)
        new_table = doc.tables[-1]
        url_paragraph = new_table.cell(0, 1).paragraphs[0]
        add_hyperlink(url_paragraph, remote_url, remote_url)
        new_table.cell(1, 1).text = commits

delete_table(doc.tables[1])

doc.save(f'./dist/{filename}')

print('File created: ', f'./dist/{filename}')

